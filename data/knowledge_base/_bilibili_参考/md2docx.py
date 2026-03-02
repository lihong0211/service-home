#!/usr/bin/env python3
"""将同目录下的指定 .md 转为 .docx（仅运行一次，可删）。"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DIR = Path(__file__).resolve().parent

def strip_md(s: str) -> str:
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = re.sub(r'^[\*\_]\s+', '', s)
    return s.strip()

def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding='utf-8')
    doc = Document()
    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(strip_md(line[2:]), level=0)
        elif line.startswith('## '):
            doc.add_heading(strip_md(line[3:]), level=1)
        elif line.startswith('### '):
            doc.add_heading(strip_md(line[4:]), level=2)
        elif line.strip() == '---':
            continue
        elif line.startswith('- '):
            doc.add_paragraph('\u2022 ' + strip_md(line[2:]))  # • 实心圆点，避免 List Bullet 渲染成空白方框
        else:
            doc.add_paragraph(strip_md(line))
    doc.save(docx_path)
    print('Written:', docx_path.name)

def main():
    # 全部 12 个 B 站知识库文档转为 docx，保留原 md
    names = [
        '哔哩哔哩公司及产品介绍',
        '哔哩哔哩内容品类与分区介绍',
        '哔哩哔哩社区公约与创作规范',
        '哔哩哔哩投稿与直播规范',
        '哔哩哔哩版权与举报申诉规则',
        '哔哩哔哩大会员权益与价格',
        '哔哩哔哩会员购与充电打赏规则',
        '哔哩哔哩广告与商业合作说明',
        '哔哩哔哩如何成为UP主与开通直播',
        '哔哩哔哩创作中心与数据说明',
        '哔哩哔哩活动与赛事规则',
        '哔哩哔哩用户反馈与申诉流程',
    ]
    for name in names:
        md_path = DIR / f'{name}.md'
        if md_path.exists():
            md_to_docx(md_path, DIR / f'{name}.docx')
        else:
            print('Skip (not found):', md_path)

if __name__ == '__main__':
    main()
