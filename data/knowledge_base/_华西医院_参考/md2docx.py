#!/usr/bin/env python3
"""将同目录下的指定 .md 转为 .docx（华西医院参考知识库）。"""
import re
from pathlib import Path

from docx import Document

DIR = Path(__file__).resolve().parent

def strip_md(s: str) -> str:
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = re.sub(r'^[\*\_]\s+', '', s)
    return s.strip()

def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding='utf-8')
    doc = Document()
    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(strip_md(line[2:]), level=0)
        elif line.startswith('## '):
            doc.add_heading(strip_md(line[3:]), level=1)
        elif line.startswith('### '):
            doc.add_heading(strip_md(line[4:]), level=2)
        elif line.strip() == '---':
            continue
        elif line.startswith('- '):
            doc.add_paragraph('\u2022 ' + strip_md(line[2:]))
        else:
            doc.add_paragraph(strip_md(line))
    doc.save(docx_path)
    print('Written:', docx_path.name)

def main():
    names = [
        '华西医院简介与院史',
        '华西医院院区分布与交通',
        '华西医院科室与专科介绍',
        '华西医院预约挂号与就诊流程',
        '华西医院门诊就医指南',
        '华西医院出入院服务',
        '华西医院医保服务',
        '华西医院特需医疗',
        '华西医院特色医疗技术与模式',
        '华西医院健康体检与健康管理',
        '华西医院患者权益与意见建议',
        '华西医院常见问题',
    ]
    for name in names:
        md_path = DIR / f'{name}.md'
        if md_path.exists():
            md_to_docx(md_path, DIR / f'{name}.docx')
        else:
            print('Skip (not found):', md_path)

if __name__ == '__main__':
    main()
